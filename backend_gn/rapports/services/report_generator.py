"""
Service de génération de rapports professionnels
Supporte PDF, Word, Excel et CSV
"""

import os
import io
from datetime import datetime
from typing import Dict, List, Any
import pandas as pd

# PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Excel
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.chart import BarChart, PieChart, LineChart, Reference

from django.conf import settings
from ..utils_export import GraphGeneratorService


class ReportGeneratorService:
    """Service principal de génération de rapports"""
    
    def __init__(self, data: Dict[str, Any], config: Dict[str, Any]):
        """
        Initialise le générateur de rapports
        
        Args:
            data: Données du rapport (statistiques, tableaux, etc.)
            config: Configuration (type, format, options, etc.)
        """
        self.data = data
        self.config = config
        self.format = config.get('format', 'pdf')
        self.type_rapport = config.get('type_rapport', 'personnalise')
        self.options = config.get('options_inclusion', {})
        
        # Créer le dossier de sortie si nécessaire
        self.output_dir = os.path.join(settings.MEDIA_ROOT, 'rapports')
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Initialiser le générateur de graphiques
        self.graph_service = GraphGeneratorService()
    
    def generate(self) -> tuple:
        """
        Génère le rapport selon le format demandé
        
        Returns:
            tuple: (chemin_fichier, nom_fichier)
        """
        generators = {
            'pdf': self._generate_pdf,
            'word': self._generate_word,
            'excel': self._generate_excel,
            'csv': self._generate_csv,
        }
        
        generator = generators.get(self.format)
        if not generator:
            raise ValueError(f"Format non supporté: {self.format}")
        
        return generator()
    
    def _generate_filename(self) -> str:
        """Génère un nom de fichier unique"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        type_rapport = self.type_rapport.replace('_', '-')
        extension = self.format if self.format != 'word' else 'docx'
        if self.format == 'excel':
            extension = 'xlsx'
        
        return f"rapport_{type_rapport}_{timestamp}.{extension}"
    
    # GÉNÉRATION PDF
    
    def _generate_pdf(self) -> tuple:
        """Génère un rapport PDF professionnel"""
        filename = self._generate_filename()
        filepath = os.path.join(self.output_dir, filename)
        
        # Créer le document PDF
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Container pour les éléments
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#185CD6'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#185CD6'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Titre
        titre = self.data.get('titre', 'Rapport Statistique')
        elements.append(Paragraph(titre, title_style))
        elements.append(Spacer(1, 0.5*cm))
        
        # Informations du rapport
        info_style = styles['Normal']
        date_debut = self.config.get('date_debut', '')
        date_fin = self.config.get('date_fin', '')
        elements.append(Paragraph(f"<b>Période:</b> {date_debut} au {date_fin}", info_style))
        elements.append(Paragraph(f"<b>Date de génération:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", info_style))
        elements.append(Spacer(1, 0.5*cm))
        
        # Résumé exécutif
        if 'resume' in self.data:
            elements.append(Paragraph("Résumé Exécutif", heading_style))
            elements.append(Paragraph(self.data['resume'], styles['Normal']))
            elements.append(Spacer(1, 0.3*cm))
        
        # Statistiques principales
        if 'statistiques' in self.data and self.options.get('includeStatistiques', True):
            elements.append(Paragraph("Statistiques Principales", heading_style))
            stats = self.data['statistiques']
            
            # Tableau de statistiques
            stats_data = [['Indicateur', 'Valeur']]
            for key, value in stats.items():
                stats_data.append([str(key), str(value)])
            
            stats_table = Table(stats_data, colWidths=[8*cm, 4*cm])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#185CD6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(stats_table)
            elements.append(Spacer(1, 0.5*cm))
        
        # Données détaillées
        if 'donnees' in self.data and self.options.get('includeDetails', False):
            elements.append(PageBreak())
            elements.append(Paragraph("Données Détaillées", heading_style))
            
            donnees = self.data['donnees']
            if donnees:
                # En-têtes
                headers = list(donnees[0].keys()) if donnees else []
                table_data = [headers]
                
                # Données
                for row in donnees[:50]:  # Limiter à 50 lignes
                    table_data.append([str(row.get(h, '')) for h in headers])
                
                details_table = Table(table_data)
                details_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#185CD6')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
                ]))
                elements.append(details_table)
        
        # Pied de page
        elements.append(Spacer(1, 1*cm))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        elements.append(Paragraph(
            f"Rapport généré par SGIC - Gendarmerie Nationale - {datetime.now().year}",
            footer_style
        ))
        
        if self.options.get('includeGraphiques', False) or self.options.get('include_graphs', False):
            try:
                rapport_id = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                graphs = self.graph_service.generate_all_graphs(self.data, rapport_id)
                
                for graph_name, graph_path in graphs.items():
                    if graph_path and os.path.exists(graph_path):
                        elements.append(Spacer(1, 0.5*cm))
                        elements.append(Paragraph(f"Graphique: {graph_name.replace('_', ' ').title()}", heading_style))
                        elements.append(Spacer(1, 0.2*cm))
                        
                        # Ajouter l'image
                        img = Image(graph_path, width=16*cm, height=9*cm)
                        elements.append(img)
                        elements.append(Spacer(1, 0.3*cm))
            except Exception as e:
                # En cas d'erreur, continuer sans graphiques
                print(f"Erreur génération graphiques PDF: {e}")
        
        # Construire le PDF
        doc.build(elements)
        
        return filepath, filename
    
    # GÉNÉRATION WORD
    
    def _generate_word(self) -> tuple:
        """Génère un rapport Word professionnel"""
        filename = self._generate_filename()
        filepath = os.path.join(self.output_dir, filename)
        
        # Créer le document
        doc = Document()
        
        # Styles personnalisés
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Titre
        titre = self.data.get('titre', 'Rapport Statistique')
        title = doc.add_heading(titre, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(24, 92, 214)  # #185CD6
        
        # Informations
        doc.add_paragraph(f"Période: {self.config.get('date_debut', '')} au {self.config.get('date_fin', '')}")
        doc.add_paragraph(f"Date de génération: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()
        
        # Résumé exécutif
        if 'resume' in self.data:
            doc.add_heading('Résumé Exécutif', 1)
            doc.add_paragraph(self.data['resume'])
        
        # Statistiques principales
        if 'statistiques' in self.data and self.options.get('includeStatistiques', True):
            doc.add_heading('Statistiques Principales', 1)
            stats = self.data['statistiques']
            
            # Tableau
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # En-têtes
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Indicateur'
            hdr_cells[1].text = 'Valeur'
            
            # Données
            for key, value in stats.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Données détaillées
        if 'donnees' in self.data and self.options.get('includeDetails', False):
            doc.add_page_break()
            doc.add_heading('Données Détaillées', 1)
            
            donnees = self.data['donnees']
            if donnees:
                headers = list(donnees[0].keys())
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # En-têtes
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                for row_data in donnees[:100]:
                    row_cells = table.add_row().cells
                    for i, header in enumerate(headers):
                        row_cells[i].text = str(row_data.get(header, ''))
        
        # Pied de page
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Rapport généré par SGIC - Gendarmerie Nationale - {datetime.now().year}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sauvegarder
        doc.save(filepath)
        
        return filepath, filename
    
    # ============================================
    # GÉNÉRATION EXCEL
    # ============================================
    
    def _generate_excel(self) -> tuple:
        """Génère un rapport Excel professionnel avec graphiques"""
        filename = self._generate_filename()
        filepath = os.path.join(self.output_dir, filename)
        
        # Créer le workbook
        wb = Workbook()
        
        # Feuille de résumé
        ws_resume = wb.active
        ws_resume.title = "Résumé"
        
        # Styles
        header_fill = PatternFill(start_color="185CD6", end_color="185CD6", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        title_font = Font(bold=True, size=16, color="185CD6")
        
        # Titre
        ws_resume['A1'] = self.data.get('titre', 'Rapport Statistique')
        ws_resume['A1'].font = title_font
        ws_resume.merge_cells('A1:D1')
        
        # Informations
        ws_resume['A3'] = 'Période:'
        ws_resume['B3'] = f"{self.config.get('date_debut', '')} au {self.config.get('date_fin', '')}"
        ws_resume['A4'] = 'Date de génération:'
        ws_resume['B4'] = datetime.now().strftime('%d/%m/%Y %H:%M')
        
        # Statistiques
        if 'statistiques' in self.data:
            ws_resume['A6'] = 'Statistiques Principales'
            ws_resume['A6'].font = Font(bold=True, size=14, color="185CD6")
            
            ws_resume['A7'] = 'Indicateur'
            ws_resume['B7'] = 'Valeur'
            ws_resume['A7'].fill = header_fill
            ws_resume['B7'].fill = header_fill
            ws_resume['A7'].font = header_font
            ws_resume['B7'].font = header_font
            
            row = 8
            for key, value in self.data['statistiques'].items():
                ws_resume[f'A{row}'] = key
                ws_resume[f'B{row}'] = value
                row += 1
        
        # Feuille de données détaillées
        if 'donnees' in self.data and self.options.get('includeDetails', False):
            ws_donnees = wb.create_sheet("Données Détaillées")
            donnees = self.data['donnees']
            
            if donnees:
                headers = list(donnees[0].keys())
                
                # En-têtes
                for col, header in enumerate(headers, 1):
                    cell = ws_donnees.cell(row=1, column=col, value=header)
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal='center')
                
                # Données
                for row_idx, row_data in enumerate(donnees, 2):
                    for col_idx, header in enumerate(headers, 1):
                        ws_donnees.cell(row=row_idx, column=col_idx, value=row_data.get(header, ''))
                
                # Ajuster la largeur des colonnes
                for col in ws_donnees.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws_donnees.column_dimensions[column].width = adjusted_width
        
        # Sauvegarder
        wb.save(filepath)
        
        return filepath, filename
    
    # GÉNÉRATION CSV
    
    def _generate_csv(self) -> tuple:
        """Génère un rapport CSV simple"""
        filename = self._generate_filename()
        filepath = os.path.join(self.output_dir, filename)
        
        # Préparer les données
        if 'donnees' in self.data and self.data['donnees']:
            df = pd.DataFrame(self.data['donnees'])
        elif 'statistiques' in self.data:
            # Si pas de données détaillées, exporter les statistiques
            df = pd.DataFrame([self.data['statistiques']])
        else:
            # Données vides
            df = pd.DataFrame([{'message': 'Aucune donnée disponible'}])
        
        # Sauvegarder en CSV
        df.to_csv(filepath, index=False, encoding='utf-8-sig', sep=';')
        
        return filepath, filename

